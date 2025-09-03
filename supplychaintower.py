# supplychaintower.py
"""
Chapter 4 – Multi-Agent Control Tower with LangChain PM Agent (LLM)
- ForecastingAgent: per-SKU HistGradientBoostingRegressor (lags + seasonal sin/cos)
- InventoryAgent: ROP + EOQ (2-week lead), ABC–XYZ service levels
- PMAgentLLM: LangChain + OpenAI (model defaults to "gpt-5"); falls back to rules if no key

Run:
  pip install -U pandas numpy scikit-learn scipy matplotlib python-docx langchain langchain-openai openai
  export OPENAI_API_KEY="sk-..." ; export OPENAI_MODEL="gpt-5"
  python chapter4_llm_pipeline.py --data beverage_sales.xlsx --outdir ./out
"""

import os, re, math, json, argparse, warnings
from collections import defaultdict
from typing import TYPE_CHECKING
import numpy as np
import pandas as pd

# Optional dependencies -------------------------------------------------------
HAS_SKLEARN = True
try:  # scikit-learn (required for GB boosting; fallback provided if missing)
    from sklearn.ensemble import HistGradientBoostingRegressor  # type: ignore
except Exception:
    HAS_SKLEARN = False

try:  # tqdm (progress bars)
    from tqdm import tqdm  # type: ignore
    HAS_TQDM = True
except Exception:
    HAS_TQDM = False

if TYPE_CHECKING:
    pass  # Skip optional imports for type checking if packages missing
import importlib

def detect_columns(df, date_col=None, sku_col=None, qty_col=None):
    """Detect date, SKU, and quantity columns with optional explicit overrides.

    Quantity detection purposely excludes price-related fields (e.g., Unit_Price, Price) unless
    explicitly specified via qty_col.
    """
    cols = list(df.columns)

    def _pick(explicit_name, candidates, purpose):
        if explicit_name is not None:
            if explicit_name not in cols:
                raise ValueError(f"{purpose} column '{explicit_name}' not found. Available: {cols}")
            return explicit_name
        if not candidates:
            raise ValueError(f"Could not detect {purpose} column. Pass --{purpose.lower()}-col explicitly. Columns: {cols}")
        return candidates[0]

    # Date candidates
    date_like = []
    for c in cols:
        lc = c.lower()
        if re.search(r'(order[\s_]*date|invoice[\s_]*date|sale[\s_]*date|^date$|datetime|time|timestamp)', lc):
            if not re.search(r'(id|number|no)$', lc):
                date_like.append(c)

    # SKU candidates
    sku_like = [c for c in cols if re.search(r'(sku|product|item|product[\s_]*id|sku[\s_]*id)', c, re.I)]

    # Quantity candidates (exclude price columns)
    qty_like = []
    for c in cols:
        lc = c.lower()
        if re.search(r'(quantity|qty|units|unit|sales[\s_]*qty|sold)', lc):
            qty_like.append(c)
    # Remove any candidate that also contains 'price'
    qty_like = [c for c in qty_like if 'price' not in c.lower()]

    dcol = _pick(date_col, date_like, 'date')
    scol = _pick(sku_col, sku_like, 'sku')
    qcol = _pick(qty_col, qty_like, 'qty')
    return dcol, scol, qcol

# LangChain availability flag (checked dynamically; avoids hard import for IDE linters)
import importlib.util as _il_util
HAS_LC = _il_util.find_spec("langchain_openai") is not None

warnings.filterwarnings("ignore")

# ---------- Data utils ----------
def load_dataset(path, date_col=None, sku_col=None, qty_col=None):
    """Load Excel/CSV dataset and return normalized DataFrame with Date, SKU, Quantity.

    Column detection is automatic but can be overridden via explicit names.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Data file not found: {path}")

    # Load
    df = pd.read_excel(path) if path.lower().endswith('.xlsx') else pd.read_csv(path)

    # Detect columns
    d, s, q = detect_columns(df, date_col=date_col, sku_col=sku_col, qty_col=qty_col)
    print(f"[load_dataset] Using columns: date={d} sku={s} qty={q}")

    # Parse dates (robust retry with dayfirst)
    dt = pd.to_datetime(df[d], errors='coerce', infer_datetime_format=True)
    if dt.isna().mean() > 0.20:
        dt2 = pd.to_datetime(df[d], errors='coerce', dayfirst=True, infer_datetime_format=True)
        if dt2.isna().mean() < dt.isna().mean():
            dt = dt2
    if dt.isna().any():
        miss_ratio = dt.isna().mean()
        if miss_ratio > 0.50:
            bad = df.loc[dt.isna(), d].astype(str).head(8).tolist()
            raise ValueError(
                f'More than 50% of date values failed to parse (ratio={miss_ratio:.2%}). Examples: ' + ', '.join(bad)
            )
        else:
            dropped = int(dt.isna().sum())
            print(f"[load_dataset] Dropping {dropped} rows with unparseable dates ({miss_ratio:.2%}).")
            df = df.loc[~dt.isna()].copy()
            dt = dt[~dt.isna()]

    out = pd.DataFrame({
        'Date': dt,
        'SKU': df[s].astype(str),
        'Quantity': pd.to_numeric(df[q], errors='coerce').fillna(0.0).astype(float)
    }).sort_values('Date').reset_index(drop=True)
    return out


def enforce_47_skus(df):
    tot = df.groupby("SKU")["Quantity"].sum().sort_values(ascending=False)
    n = len(tot)
    if n > 47:
        keep = tot.head(47).index
        return df[df["SKU"].isin(keep)].copy(), f"Dataset had {n} SKUs; kept top 47 by volume."
    elif n < 47:
        return df.copy(), f"Dataset has {n} SKUs (<47). Proceeding."
    return df.copy(), "Dataset has 47 SKUs."

def weekly_aggregate(df):
    start = df["Date"].min()
    df["DayIndex"] = (df["Date"] - start).dt.days
    df["WeekIndex"] = df["DayIndex"] // 7
    maxw = int(df["WeekIndex"].max())
    if ((df["DayIndex"] // 7) == maxw).sum() < 7:
        df = df[df["WeekIndex"] < maxw].copy()
        maxw -= 1
    weekly = df.groupby(["SKU","WeekIndex"], as_index=False).agg(WeeklySales=("Quantity","sum"))
    weekly["Week"] = weekly["WeekIndex"] + 1
    all_skus = weekly["SKU"].unique()
    full = pd.MultiIndex.from_product([all_skus, range(0, maxw+1)], names=["SKU","WeekIndex"])
    weekly = weekly.set_index(["SKU","WeekIndex"]).reindex(full, fill_value=0.0).reset_index()
    weekly["Week"] = weekly["WeekIndex"] + 1
    weekly = weekly[weekly["Week"] <= 52].copy()
    return weekly

def abc_xyz_segmentation(weekly, prices=None):
    skus = weekly["SKU"].unique()
    if prices is None:
        rng = np.random.default_rng(0)
        prices = {sku: float(rng.uniform(5,20)) for sku in skus}
    ann = weekly.groupby("SKU")["WeeklySales"].sum().reset_index()
    ann["AnnualValue"] = ann.apply(lambda r: r["WeeklySales"]*prices[r["SKU"]], axis=1)
    ann = ann.sort_values("AnnualValue", ascending=False).reset_index(drop=True)
    totv = ann["AnnualValue"].sum()
    ann["CumPct"] = ann["AnnualValue"].cumsum() / totv * 100 if totv>0 else 0
    def abc(p): return "A" if p<=80 else ("B" if p<=95 else "C")
    ann["ABC"] = ann["CumPct"].apply(abc)

    g = weekly.groupby("SKU")["WeeklySales"].agg(["mean","std"]).reset_index().fillna(0.0)
    g["CV"] = g.apply(lambda r: (r["std"]/r["mean"]) if r["mean"]>0 else 10.0, axis=1)
    g.loc[~np.isfinite(g["CV"]), "CV"] = 10.0
    def xyz(cv): return "X" if cv<=0.5 else ("Y" if cv<=1.0 else "Z")
    g["XYZ"] = g["CV"].apply(xyz)

    seg = ann[["SKU","ABC"]].merge(g[["SKU","XYZ","CV","mean","std"]], on="SKU", how="left")
    seg["Segment"] = seg["ABC"] + "-" + seg["XYZ"]
    return seg, prices

def z_from_service(p):
    table = {0.75:0.674, 0.80:0.842, 0.82:0.915, 0.85:1.036, 0.88:1.175, 0.90:1.282, 0.92:1.405,
             0.95:1.645, 0.98:2.054, 0.99:2.326, 0.995:2.576}
    ks = np.array(list(table.keys()))
    return table[float(ks[np.argmin(np.abs(ks - p))])]

# ---------- Agents ----------
class ForecastingAgent:
    """Per-SKU forecasting with optional gradient boosting.

    Falls back to a simple moving-average style regressor if scikit-learn isn't available.
    """

    class _NaiveRegressor:
        def fit(self, X, y):
            self.value = float(np.mean(y[-3:])) if len(y) else 0.0
        def predict(self, X):  # shape (n, features)
            return np.full((len(X),), self.value, dtype=float)

    def __init__(self):
        self.models, self.last_two = {}, {}

    def fit(self, weekly, train_weeks=44):
        for sku, d in weekly[weekly["Week"] <= train_weeks].groupby("SKU"):
            d = d.sort_values("Week").copy()
            d["Lag1"] = d["WeeklySales"].shift(1)
            d["Lag2"] = d["WeeklySales"].shift(2)
            d["Week_sin"] = np.sin(2*np.pi*d["Week"]/52.0)
            d["Week_cos"] = np.cos(2*np.pi*d["Week"]/52.0)
            d["Trend"] = d["Week"]
            d = d.dropna()
            if len(d) < 6:
                continue
            X = d[["Lag1","Lag2","Week_sin","Week_cos","Trend"]].values
            y = d["WeeklySales"].values
            if HAS_SKLEARN:
                model = HistGradientBoostingRegressor(max_depth=3, max_iter=200, learning_rate=0.05, random_state=0)
            else:
                model = self._NaiveRegressor()
            model.fit(X, y)
            self.models[sku] = model
            tail = weekly[(weekly["SKU"]==sku) & (weekly["Week"]<=train_weeks)].sort_values("Week")["WeeklySales"].tail(2).tolist()
            self.last_two[sku] = tail if len(tail) == 2 else ([tail[0], tail[0]] if tail else [0, 0])

    def forecast(self, sku, week):
        p1, p2 = self.last_two.get(sku, [0, 0])
        X = np.array([[p1, p2, math.sin(2*math.pi*week/52.0), math.cos(2*math.pi*week/52.0), week]])
        yhat = float(self.models[sku].predict(X)[0]) if sku in self.models else p1
        return max(0.0, yhat)

    def update_last_two(self, sku, actual):
        self.last_two[sku] = [self.last_two.get(sku, [actual, actual])[-1], actual]

class InventoryAgent:
    def __init__(self, seg_df, prices):
        self.seg_map = seg_df.set_index("SKU")["Segment"].to_dict()
        self.prices = prices
        self.service = {"A-X":0.99,"A-Y":0.95,"A-Z":0.98,"B-X":0.92,"B-Y":0.91,"B-Z":0.88,"C-X":0.85,"C-Y":0.82,"C-Z":0.75}
        self.order_cost, self.holding_rate, self.L = 100.0, 0.20, 2
        self.rop, self.eoq = {}, {}
        self.on_hand, self.on_order, self.pipeline = defaultdict(float), defaultdict(float), defaultdict(list)
        self.mu, self.sd = defaultdict(float), defaultdict(float)
    def initialize(self, weekly, train_weeks=44):
        s = weekly[weekly["Week"]<=train_weeks].groupby("SKU")["WeeklySales"].agg(["mean","std"]).fillna(0.0)
        for sku in s.index:
            self.mu[sku], self.sd[sku] = float(s.loc[sku,"mean"]), float(s.loc[sku,"std"])
            seg, sl = self.seg_map.get(sku,"C-Z"), self.service.get(self.seg_map.get(sku,"C-Z"),0.85)
            z, mu, sd = z_from_service(sl), self.mu[sku], self.sd[sku]
            rop = self.L*mu + z*sd*math.sqrt(self.L)
            D = mu*52.0; H = self.holding_rate*self.prices.get(sku,10.0)
            eoq = math.sqrt((2*max(D,0.0)*self.order_cost)/max(H,1e-6)) if D>0 else 0.0
            self.rop[sku], self.eoq[sku] = math.ceil(rop), (max(1.0, math.floor(min(eoq,D))) if eoq>0 else 0.0)
            self.on_hand[sku] = self.rop[sku] + math.ceil(mu); self.on_order[sku] = 0.0; self.pipeline[sku] = []
    def process_arrivals(self, sku):
        new = []
        for t,q in self.pipeline[sku]:
            if t<=0: self.on_hand[sku]+=q; self.on_order[sku]-=q
            else: new.append((t-1,q))
        self.pipeline[sku]=new
    def recompute_rop_with_service(self, sku, new_service):
        z = z_from_service(new_service); mu, sd = self.mu.get(sku,0.0), self.sd.get(sku,0.0)
        self.rop[sku] = math.ceil(self.L*mu + z*sd*math.sqrt(self.L))
    def review_and_order(self, sku, early_trigger=False, trigger_forecast=0.0, relax_cap=1.0):
        inv_pos, trigger = self.on_hand[sku]+self.on_order[sku], self.rop.get(sku,0.0)
        if early_trigger: trigger += trigger_forecast
        if inv_pos <= trigger:
            q = self.eoq.get(sku,0.0)
            if q>0:
                q_eff = q*float(max(1.0,relax_cap))
                self.pipeline[sku].append((self.L,q_eff)); self.on_order[sku]+=q_eff
                return q_eff
        return 0.0

# ---------- PM Agent (LangChain/OpenAI) ----------
PM_PROMPT = """You are a Project Manager AI for beverage replenishment.
Return STRICT JSON:
- new_service: float in [0.75, 0.995]
- early_trigger: boolean
- relax_cap: float >= 1.0
- uplift_pct: float in [0.0, 0.3]
Context:
- segment: {segment}
- previous_service: {prev_service}
- forecast_this_week: {forecast_this_week}
- inventory_position: {inventory_position}
- rop_current: {rop_current}
- eoq: {eoq}
- mean_weekly_demand: {mu}
- std_weekly_demand: {sd}
- last_week_stockout: {last_week_stockout}
Principles:
- A-class prioritizes service; C-class prioritizes efficiency.
- Raise service for A/B if shortfall risk (inv_pos < LT demand).
- Reduce service for C if cover excessive.
- relax_cap >1.0 mainly for A/AZ; uplift_pct conservative (no external signals).
Return JSON only.
"""

class PMAgentLLM:
    def __init__(self, inv_agent):
        self.inv = inv_agent
        self.has_key = bool(os.getenv("OPENAI_API_KEY"))
        self.model_name = os.getenv("OPENAI_MODEL","gpt-5-nano-2025-08-07")
        self.llm = None
        if HAS_LC and self.has_key and importlib.util.find_spec("langchain_openai") is not None:
            try:
                    from langchain_openai import ChatOpenAI  # type: ignore  # optional
                    self.llm = ChatOpenAI(model=self.model_name, temperature=0.1)  # pragma: no cover
            except Exception as e:
                print(f"[PM LLM] Init error: {e}"); self.llm = None
    def _rule_based(self, sku, forecast, last_week_stockout):
        seg, ABC = self.inv.seg_map.get(sku,"C-Z"), self.inv.seg_map.get(sku,"C-Z")[0]
        base = self.inv.service.get(seg,0.85)
        early, relax, new_service = False, 1.0, base
        inv_pos = self.inv.on_hand[sku]+self.inv.on_order[sku]
        if ABC in ("A","B") and (inv_pos < forecast*self.inv.L or last_week_stockout):
            new_service = min(0.995, base+0.02); early=True; relax=1.5 if seg=="A-Z" else 1.2
        if ABC=="C" and self.inv.on_hand[sku] > 6*max(1.0, forecast):
            new_service = max(0.75, base-0.05)
        return {"new_service":new_service,"early_trigger":early,"relax_cap":relax,"uplift_pct":0.0}
    def adjust_policy(self, sku, week, forecast_this_week, last_week_stockout=False):
        if self.llm is None:
            return self._rule_based(sku, forecast_this_week, last_week_stockout)
        seg = self.inv.seg_map.get(sku,"C-Z"); base = self.inv.service.get(seg,0.85)
        inv_pos = self.inv.on_hand[sku]+self.inv.on_order[sku]
        prompt = PM_PROMPT.format(segment=seg, prev_service=base, forecast_this_week=round(float(forecast_this_week),3),
                                  inventory_position=round(float(inv_pos),3), rop_current=self.inv.rop.get(sku,0.0),
                                  eoq=self.inv.eoq.get(sku,0.0), mu=round(float(self.inv.mu.get(sku,0.0)),3),
                                  sd=round(float(self.inv.sd.get(sku,0.0)),3),
                                  last_week_stockout=str(bool(last_week_stockout)).lower())
        try:
            resp = self.llm.invoke(prompt).content
            m = re.search(r"\{.*\}", resp, flags=re.S)
            if not m: raise ValueError("No JSON found in LLM response.")
            data = json.loads(m.group(0))
            new_service = float(min(0.995, max(0.75, data.get("new_service", base))))
            early_trigger = bool(data.get("early_trigger", False))
            relax_cap = float(max(1.0, data.get("relax_cap", 1.0)))
            uplift_pct = float(min(0.3, max(0.0, data.get("uplift_pct", 0.0))))
            return {"new_service":new_service,"early_trigger":early_trigger,"relax_cap":relax_cap,"uplift_pct":uplift_pct}
        except Exception:
            return self._rule_based(sku, forecast_this_week, last_week_stockout)

# ---------- KPIs ----------
def kpis(sim_df, test_weeks):
    df = sim_df.copy()
    df["APE"] = np.where(df["Actual"]>0, np.abs(df["Forecast"]-df["Actual"])/df["Actual"], np.nan)
    mape = float(np.nanmean(df["APE"]))
    fill = float(df["Fulfilled"].sum()/max(df["Actual"].sum(),1e-6))
    avg_inv = float(df["EndingStock"].mean())
    annual_demand = df["Actual"].sum()*(52.0/len(test_weeks))
    turn = float(annual_demand/max(avg_inv,1e-6))
    so_weeks = int((df["StockoutFlag"]>0).sum())
    hold = float(df["HoldingCost"].sum()); order = float(df["OrderingCost"].sum())
    return {"MAPE":mape,"FillRate":fill,"AvgInventory":avg_inv,"Turnover":turn,
            "StockoutWeeks":so_weeks,"HoldingCost":hold,"OrderingCost":order,"TotalCost":hold+order}

# ---------- Main ----------
def main(args):
    data = load_dataset(args.data, date_col=args.date_col, sku_col=args.sku_col, qty_col=args.qty_col)
    data, note = enforce_47_skus(data); print("SKU check:", note)
    weekly = weekly_aggregate(data)
    total_weeks = weekly["Week"].nunique()
    train_weeks = min(44, total_weeks-8)
    test_weeks = list(range(train_weeks+1, total_weeks+1))
    train = weekly[weekly["Week"]<=train_weeks].copy()
    test  = weekly[weekly["Week"]>train_weeks].copy()
    print(f"Weeks: total={total_weeks}, train=1..{train_weeks}, test={test_weeks[0]}..{test_weeks[-1]}")

    seg, prices = abc_xyz_segmentation(weekly)

    # Forecasting
    f_agent = ForecastingAgent(); f_agent.fit(weekly, train_weeks=train_weeks)

    # Scenarios: Baseline (no PM), PM_Lean (LLM), PM_HighService (also LLM; A gets higher base service via decisions)
    scenarios = {"Baseline":{}, "PM_Lean":{}, "PM_HighService":{}}
    results = {}

    # ---------------- Progress Configuration -----------------
    # Extended progress display with mode selection
    progress_mode = getattr(args, 'progress_mode', 'detailed') if hasattr(args, 'progress_mode') else 'detailed'
    use_progress = (not getattr(args, 'no_progress', False)) and HAS_TQDM and progress_mode != 'off'
    sku_list = sorted(weekly["SKU"].unique())

    def make_bar(name):
        if not use_progress:
            return None
        if progress_mode == 'simple':
            return tqdm(total=len(test_weeks), desc=f"{name}", unit="week", leave=True)
        elif progress_mode == 'detailed':
            total = len(test_weeks) * len(sku_list)
            return tqdm(total=total, desc=f"{name} W?", unit="sku-week", leave=True)
        return None

    for scen in scenarios:
        bar = make_bar(scen)
        inv = InventoryAgent(seg, prices); inv.initialize(weekly, train_weeks=train_weeks)
        pm  = PMAgentLLM(inv)
        sim_rows, last_so = [], defaultdict(lambda: False)

        for wk in test_weeks:
            # Step 1: forecasts
            forecasts = {sku: f_agent.forecast(sku, wk) for sku in sku_list}
            # Step 2: actual demand for this week
            dm = {r["SKU"]: float(r["WeeklySales"]) for _, r in test[test["Week"] == wk].iterrows()}

            if bar and progress_mode == 'detailed':
                bar.set_description(f"{scen} W{wk}")

            for sku in sku_list:
                # Arrivals / pipeline decrement
                inv.process_arrivals(sku)
                # PM decision (dynamic except baseline)
                if scen == "Baseline":
                    seg_code = inv.seg_map.get(sku, "C-Z")
                    base_service = inv.service.get(seg_code, 0.85)
                    dec = {"new_service": base_service, "early_trigger": False, "relax_cap": 1.0, "uplift_pct": 0.0}
                else:
                    dec = pm.adjust_policy(sku, wk, forecasts[sku], last_so[sku])
                    inv.recompute_rop_with_service(sku, dec["new_service"])
                trigger_forecast = forecasts[sku] * (1.0 + dec.get("uplift_pct", 0.0))
                # Demand realization
                actual = dm.get(sku, 0.0)
                begin = inv.on_hand[sku]
                shipped = min(begin, actual)
                stockout = 1 if shipped < actual else 0
                inv.on_hand[sku] = max(0.0, begin - shipped)
                end = inv.on_hand[sku]; last_so[sku] = (stockout == 1)
                # Order decision
                prev_on_order = inv.on_order[sku]
                inv.review_and_order(sku, early_trigger=dec["early_trigger"], trigger_forecast=trigger_forecast, relax_cap=dec["relax_cap"])
                ordered_now = inv.on_order[sku] - prev_on_order if inv.on_order[sku] > prev_on_order else 0.0
                # Costs
                avg_inv = (begin + end) / 2.0; unit_cost = prices.get(sku, 10.0)
                hold = avg_inv * unit_cost * (inv.holding_rate / 52.0)
                order_cost = inv.order_cost if ordered_now > 0 else 0.0
                sim_rows.append({
                    "Scenario": scen, "Week": wk, "SKU": sku, "Forecast": forecasts[sku], "Actual": actual,
                    "BeginningStock": begin, "EndingStock": end, "Fulfilled": shipped, "LostSales": actual - shipped,
                    "StockoutFlag": stockout, "OrderPlacedQty": ordered_now, "HoldingCost": hold, "OrderingCost": order_cost,
                    "PM_new_service": dec["new_service"], "PM_early_trigger": dec["early_trigger"],
                    "PM_relax_cap": dec["relax_cap"], "PM_uplift_pct": dec["uplift_pct"]
                })
                f_agent.update_last_two(sku, actual)
                if bar and progress_mode == 'detailed':
                    bar.update(1)

            if bar and progress_mode == 'simple':
                bar.update(1)
                bar.set_postfix({"week": wk, "scen": scen})
            elif bar and progress_mode == 'detailed':
                # periodic postfix update at end of week
                bar.set_postfix({"week": wk})

        if bar:
            bar.close()

        sim_df = pd.DataFrame(sim_rows); results[scen] = {"sim": sim_df, "kpis": kpis(sim_df, test_weeks)}

    # KPI table
    kpi_table = pd.DataFrame([
        {"Metric":"MAPE","Baseline":results["Baseline"]["kpis"]["MAPE"],"PM_Lean":results["PM_Lean"]["kpis"]["MAPE"],"PM_HighService":results["PM_HighService"]["kpis"]["MAPE"]},
        {"Metric":"Fill Rate","Baseline":results["Baseline"]["kpis"]["FillRate"],"PM_Lean":results["PM_Lean"]["kpis"]["FillRate"],"PM_HighService":results["PM_HighService"]["kpis"]["FillRate"]},
        {"Metric":"Avg Inventory (units)","Baseline":results["Baseline"]["kpis"]["AvgInventory"],"PM_Lean":results["PM_Lean"]["kpis"]["AvgInventory"],"PM_HighService":results["PM_HighService"]["kpis"]["AvgInventory"]},
        {"Metric":"Turnover","Baseline":results["Baseline"]["kpis"]["Turnover"],"PM_Lean":results["PM_Lean"]["kpis"]["Turnover"],"PM_HighService":results["PM_HighService"]["kpis"]["Turnover"]},
        {"Metric":"Stockout Weeks","Baseline":results["Baseline"]["kpis"]["StockoutWeeks"],"PM_Lean":results["PM_Lean"]["kpis"]["StockoutWeeks"],"PM_HighService":results["PM_HighService"]["kpis"]["StockoutWeeks"]},
        {"Metric":"Holding Cost","Baseline":results["Baseline"]["kpis"]["HoldingCost"],"PM_Lean":results["PM_Lean"]["kpis"]["HoldingCost"],"PM_HighService":results["PM_HighService"]["kpis"]["HoldingCost"]},
        {"Metric":"Ordering Cost","Baseline":results["Baseline"]["kpis"]["OrderingCost"],"PM_Lean":results["PM_Lean"]["kpis"]["OrderingCost"],"PM_HighService":results["PM_HighService"]["kpis"]["OrderingCost"]},
        {"Metric":"Total Cost","Baseline":results["Baseline"]["kpis"]["TotalCost"],"PM_Lean":results["PM_Lean"]["kpis"]["TotalCost"],"PM_HighService":results["PM_HighService"]["kpis"]["TotalCost"]},
    ])

    # Segment performance (PM_HighService)
    seg_perf = results["PM_HighService"]["sim"].merge(seg[["SKU","Segment"]], on="SKU", how="left").groupby("Segment").agg(
        SKUs=("SKU","nunique"),
        Demand=("Actual","sum"),
        Shipped=("Fulfilled","sum"),
        FillRate=("Fulfilled", lambda s: s.sum()/results["PM_HighService"]["sim"].loc[s.index,"Actual"].sum() if results["PM_HighService"]["sim"].loc[s.index,"Actual"].sum()>0 else 1.0),
        AvgOnHand=("EndingStock","mean"),
        StockoutWeeks=("StockoutFlag","sum"),
        Orders=("OrderPlacedQty", lambda q: (q>0).sum())
    ).reset_index()

    # Save
    os.makedirs(args.outdir, exist_ok=True)
    kpi_table.to_csv(os.path.join(args.outdir,"kpis_summary.csv"), index=False)
    seg_perf.to_csv(os.path.join(args.outdir,"segment_performance.csv"), index=False)
    for name, bundle in results.items():
        bundle["sim"].to_csv(os.path.join(args.outdir,f"sim_details_{name}.csv"), index=False)

    # Word tables (Elsevier style)
    try:
        import importlib.util
        if importlib.util.find_spec("docx") is None:
            raise ImportError("docx not installed")
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_heading("Chapter 4 – Experimental Results (Beverage Sales, 47 SKUs)", level=1)
        p = doc.add_paragraph("We evaluate a multi-agent control tower on the Beverage Sales dataset. ")
        p.add_run("External signals excluded; ABC–XYZ segmentation; lead time=2 weeks. ")
        p.add_run("We compare Baseline and LangChain-powered PM modes.")
        doc.add_heading("Table 1 – KPI Summary", level=2)
        t1 = doc.add_table(rows=1, cols=len(kpi_table.columns))
        for j,c in enumerate(kpi_table.columns): t1.rows[0].cells[j].text = str(c)
        for _,r in kpi_table.round(4).iterrows():
            row = t1.add_row().cells
            for j,c in enumerate(kpi_table.columns): row[j].text = str(r[c])
        doc.add_heading("Table 2 – Segment Performance (PM High Service)", level=2)
        t2 = doc.add_table(rows=1, cols=len(seg_perf.columns))
        for j,c in enumerate(seg_perf.columns): t2.rows[0].cells[j].text = str(c)
        for _,r in seg_perf.round(4).iterrows():
            row = t2.add_row().cells
            for j,c in enumerate(seg_perf.columns): row[j].text = str(r[c])
        doc.save(os.path.join(args.outdir,"Results.docx"))
    except Exception as e:
        print("Word doc skipped:", e)

    # Plot for a top A SKU
    try:
        import matplotlib.pyplot as plt  # type: ignore
        topA = seg[seg["ABC"] == "A"]["SKU"].iloc[0]
        hist = weekly[weekly["SKU"] == topA].sort_values("Week")
        fig = plt.figure()
        plt.plot(hist["Week"], hist["WeeklySales"], label="Actual")
        f = results["PM_HighService"]["sim"]
        f = f[f["SKU"] == topA][["Week", "Forecast"]]
        if not f.empty:
            plt.plot(f["Week"], f["Forecast"], label="Forecast")
        plt.title("Top A SKU — Weekly Sales vs Forecast"); plt.xlabel("Week"); plt.ylabel("Units"); plt.legend()
        plt.savefig(os.path.join(args.outdir, "plot_topA_forecast.png"), bbox_inches="tight"); plt.close(fig)
    except Exception as e:
        print("Plot skipped:", e)

    
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Multi-Agent Supply Chain Control Tower Simulation")
    parser.add_argument("--data", type=str, default="beverage_sales.csv", help="Path to Excel/CSV beverage dataset")
    parser.add_argument("--outdir", type=str, default="./out", help="Output directory for results")
    parser.add_argument("--date-col", type=str, default=None, help="Explicit date column override")
    parser.add_argument("--sku-col", type=str, default=None, help="Explicit SKU/product column override")
    parser.add_argument("--qty-col", type=str, default=None, help="Explicit quantity column override")
    parser.add_argument("--no-progress", action="store_true", help="Disable progress bars (tqdm)")
    parser.add_argument("--progress-mode", choices=["simple","detailed","off"], default="detailed", help="Progress display detail level")
    args = parser.parse_args()

    # Auto-load OpenAI key from file if present and env var not set (avoid printing key)
    if not os.getenv("OPENAI_API_KEY") and os.path.exists("openai_api.txt"):
        try:
            with open("openai_api.txt","r",encoding="utf-8") as fh:
                key = fh.read().strip()
            if key:
                os.environ["OPENAI_API_KEY"] = key
        except Exception:
            pass

    main(args)
